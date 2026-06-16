"""
doc-parser Lambda
Step 1 of the translation workflow:
  - Downloads the .docx from S3
  - Computes SHA-256 for idempotency (same file → same job_id, skip if already SUCCEEDED)
  - Parses all text chunks (paragraphs, table cells, headers, footers)
  - Detects source language via CJK character ratio
  - Writes chunk manifest to S3 and records job in DynamoDB
"""
import hashlib
import json
import os
import re
import tempfile
from datetime import datetime, timedelta, timezone

import boto3

s3 = boto3.client("s3")
dynamodb = boto3.resource("dynamodb")

CHUNKS_BUCKET = os.environ["CHUNKS_BUCKET"]
JOBS_TABLE = os.environ["JOBS_TABLE"]

# CJK Unified Ideographs block — more reliable than langdetect on mixed-language geo docs
CJK_RE = re.compile(r"[一-鿿]")


def detect_language(chunks: list[dict]) -> str:
    """Return 'zh' if CJK chars exceed 15% of the first-20-chunk sample, else 'en'."""
    sample = "".join(c["text"] for c in chunks[:20])
    if not sample:
        return "en"
    ratio = len(CJK_RE.findall(sample)) / len(sample)
    return "zh" if ratio > 0.15 else "en"


def parse_docx(path: str) -> list[dict]:
    from docx import Document  # python-docx provided by Lambda layer

    doc = Document(path)
    chunks: list[dict] = []
    idx = 0

    # Body paragraphs — ref uses para_idx (position in doc.paragraphs, including empty ones)
    # so it matches enumerate(doc.paragraphs) in assemble_docx exactly.
    for para_idx, para in enumerate(doc.paragraphs):
        if para.text.strip():
            chunks.append({
                "id": idx,
                "type": "paragraph",
                "ref": f"paragraph_{para_idx}",
                "text": para.text,
                "style": para.style.name,
            })
            idx += 1

    # Table cells
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                if cell.text.strip():
                    chunks.append({
                        "id": idx,
                        "type": "table_cell",
                        "ref": f"table_{t_idx}_{r_idx}_{c_idx}",
                        "text": cell.text,
                    })
                    idx += 1

    # Headers and footers
    for s_idx, section in enumerate(doc.sections):
        for loc, hf_obj in [("header", section.header), ("footer", section.footer)]:
            for p_idx, para in enumerate(hf_obj.paragraphs):
                if para.text.strip():
                    chunks.append({
                        "id": idx,
                        "type": loc,
                        "ref": f"{loc}_{s_idx}_{p_idx}",
                        "text": para.text,
                    })
                    idx += 1

    return chunks


def lambda_handler(event: dict, context) -> dict:
    bucket = event["bucket"]
    key = event["key"]

    # Download to temp file
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        s3.download_fileobj(bucket, key, tmp)
        tmp_path = tmp.name

    try:
        # SHA-256 → first 16 hex chars as job_id (enough entropy, shorter keys)
        with open(tmp_path, "rb") as f:
            file_bytes = f.read()
        job_id = hashlib.sha256(file_bytes).hexdigest()[:16]

        # Idempotency check
        table = dynamodb.Table(JOBS_TABLE)
        existing = table.get_item(Key={"job_id": job_id}).get("Item")
        if existing and existing.get("status") == "SUCCEEDED":
            return {
                "job_id": job_id,
                "source_lang": existing["source_lang"],
                "chunks_bucket": CHUNKS_BUCKET,
                "chunks_manifest": existing["chunks_manifest"],
                "total_chunks": int(existing["total_chunks"]),
                "source_bucket": bucket,
                "source_key": key,
                "skipped": True,
            }

        # Parse document
        chunks = parse_docx(tmp_path)
        source_lang = detect_language(chunks)

        # Embed job context into every chunk — chunk-translator needs these
        for chunk in chunks:
            chunk["job_id"] = job_id
            chunk["source_lang"] = source_lang

        # Write manifest to S3
        manifest_key = f"chunks/{job_id}/manifest.json"
        s3.put_object(
            Bucket=CHUNKS_BUCKET,
            Key=manifest_key,
            Body=json.dumps(chunks, ensure_ascii=False),
            ContentType="application/json",
        )

        # Record job; ttl = 30 days so DynamoDB auto-expires completed jobs
        now = datetime.now(timezone.utc)
        ttl = int((now + timedelta(days=30)).timestamp())
        table.put_item(Item={
            "job_id": job_id,
            "status": "PROCESSING",
            "source_lang": source_lang,
            "source_bucket": bucket,
            "source_key": key,
            "chunks_manifest": manifest_key,
            "total_chunks": len(chunks),
            "done_chunks": 0,
            "created_at": now.isoformat(),
            "ttl": ttl,
        })

        return {
            "job_id": job_id,
            "source_lang": source_lang,
            "chunks_bucket": CHUNKS_BUCKET,
            "chunks_manifest": manifest_key,
            "total_chunks": len(chunks),
            "source_bucket": bucket,
            "source_key": key,
        }
    finally:
        os.unlink(tmp_path)
