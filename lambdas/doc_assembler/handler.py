"""
doc-assembler Lambda
Step 4 of the translation workflow:
  - Downloads the original .docx (template for styles)
  - Reads all per-chunk result JSONs from S3 results/job_id/
  - Rebuilds the docx in-place using paragraph-level format preservation
    (inline/run-level mixed formatting is intentionally not preserved — see docs/architecture.md §4.5)
  - Uploads the translated .docx to the output bucket
  - Marks the job SUCCEEDED in DynamoDB
"""
import json
import logging
import os
import tempfile
from datetime import datetime, timezone

import boto3

logger = logging.getLogger()
logger.setLevel(logging.INFO)

s3 = boto3.client("s3")
dynamodb = boto3.resource("dynamodb")

RESULTS_BUCKET = os.environ["RESULTS_BUCKET"]
OUTPUT_BUCKET = os.environ["OUTPUT_BUCKET"]
JOBS_TABLE = os.environ["JOBS_TABLE"]


# ── Format-safe text replacement ─────────────────────────────────────────────

def _set_paragraph_text(para, text: str) -> None:
    """
    Write translated text into a paragraph while keeping the paragraph's
    primary run format (font, size, bold-at-paragraph-level, etc.).

    Design constraint: we only guarantee paragraph-level format fidelity.
    Inline (run-level) mixed formatting — e.g. a single bold word inside a
    normal sentence — is lost, because the LLM returns a single plain string
    and we have no reliable way to re-align run boundaries across languages.
    """
    if not text:
        return
    if para.runs:
        # Keep first run's format as the paragraph "style carrier"; clear the rest
        para.runs[0].text = text
        for run in para.runs[1:]:
            run.text = ""
    else:
        # Paragraph has no runs (pure style paragraph) — create one
        para.add_run(text)


# ── Result loading ────────────────────────────────────────────────────────────

def load_results(job_id: str) -> dict[str, dict]:
    """List and read all result JSON files for this job from S3."""
    results: dict[str, dict] = {}
    paginator = s3.get_paginator("list_objects_v2")
    prefix = f"results/{job_id}/"
    page_count = 0
    for page in paginator.paginate(Bucket=RESULTS_BUCKET, Prefix=prefix):
        for obj in page.get("Contents", []):
            body = s3.get_object(Bucket=RESULTS_BUCKET, Key=obj["Key"])["Body"].read()
            item = json.loads(body)
            results[item["ref"]] = item
        page_count += 1
    logger.info("Loaded %d translated chunks from %d S3 pages (job=%s)", len(results), page_count, job_id)
    return results


# ── Document assembly ─────────────────────────────────────────────────────────

def assemble_docx(original_path: str, results: dict[str, dict]):
    from docx import Document  # python-docx provided by Lambda layer

    doc = Document(original_path)  # open original as style template

    # Body paragraphs
    for i, para in enumerate(doc.paragraphs):
        chunk = results.get(f"paragraph_{i}")
        if chunk:
            _set_paragraph_text(para, chunk["translated"])

    # Table cells — traverse every paragraph inside each cell
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                chunk = results.get(f"table_{t_idx}_{r_idx}_{c_idx}")
                if chunk and cell.paragraphs:
                    _set_paragraph_text(cell.paragraphs[0], chunk["translated"])
                    # Clear extra paragraphs inside the cell
                    for extra in cell.paragraphs[1:]:
                        _set_paragraph_text(extra, "")

    # Headers and footers
    for s_idx, section in enumerate(doc.sections):
        for loc, hf_obj in [("header", section.header), ("footer", section.footer)]:
            for p_idx, para in enumerate(hf_obj.paragraphs):
                chunk = results.get(f"{loc}_{s_idx}_{p_idx}")
                if chunk:
                    _set_paragraph_text(para, chunk["translated"])

    return doc


# ── Lambda handler ────────────────────────────────────────────────────────────

def lambda_handler(event: dict, context) -> dict:
    job_id: str = event["job_id"]
    source_bucket: str = event["source_bucket"]
    source_key: str = event["source_key"]

    logger.info("Assembling job=%s from s3://%s/%s", job_id, source_bucket, source_key)

    # Download original docx (keep as style template)
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        s3.download_fileobj(source_bucket, source_key, tmp)
        original_path = tmp.name

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as out_tmp:
        out_path = out_tmp.name

    try:
        results = load_results(job_id)
        doc = assemble_docx(original_path, results)
        doc.save(out_path)

        # Derive output key: keep original filename, add _translated suffix
        base_name = source_key.rsplit("/", 1)[-1].replace(".docx", "")
        output_key = f"translated/{job_id}/{base_name}_translated.docx"

        with open(out_path, "rb") as f:
            s3.put_object(
                Bucket=OUTPUT_BUCKET,
                Key=output_key,
                Body=f.read(),
                ContentType=(
                    "application/vnd.openxmlformats-officedocument"
                    ".wordprocessingml.document"
                ),
            )
        logger.info("Uploaded translated doc to s3://%s/%s", OUTPUT_BUCKET, output_key)

        # Mark job complete
        dynamodb.Table(JOBS_TABLE).update_item(
            Key={"job_id": job_id},
            UpdateExpression="SET #s = :s, output_key = :k, finished_at = :t",
            ExpressionAttributeNames={"#s": "status"},
            ExpressionAttributeValues={
                ":s": "SUCCEEDED",
                ":k": output_key,
                ":t": datetime.now(timezone.utc).isoformat(),
            },
        )

        return {
            "job_id": job_id,
            "output_bucket": OUTPUT_BUCKET,
            "output_key": output_key,
            "status": "SUCCEEDED",
            "chunks_assembled": len(results),
        }
    finally:
        os.unlink(original_path)
        os.unlink(out_path)
